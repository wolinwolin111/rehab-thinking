from pathlib import Path
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path(__file__).resolve().parents[1] / "docs" / "rehabmind-pilot-startup-plan.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
MUTED = "5F6B7A"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"
CAUTION = "FFF8E8"
RISK = "FDECEC"
WHITE = "FFFFFF"

CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_border(cell, color="D6DEE8", size="6"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = qn(f"w:{edge}")
        element = borders.find(tag)
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_paragraph_box(paragraph, fill=None, border_color="D6DEE8", border_size="6", spacing=100):
    p_pr = paragraph._p.get_or_add_pPr()
    if fill:
        shd = p_pr.find(qn("w:shd"))
        if shd is None:
            shd = OxmlElement("w:shd")
            p_pr.append(shd)
        shd.set(qn("w:fill"), fill)
        shd.set(qn("w:val"), "clear")
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    for edge in ("top", "left", "bottom", "right"):
        tag = qn(f"w:{edge}")
        element = p_bdr.find(tag)
        if element is None:
            element = OxmlElement(f"w:{edge}")
            p_bdr.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), border_size)
        element.set(qn("w:space"), "4")
        element.set(qn("w:color"), border_color)
    paragraph.paragraph_format.left_indent = Pt(6)
    paragraph.paragraph_format.right_indent = Pt(6)
    paragraph.paragraph_format.space_before = Pt(3)
    paragraph.paragraph_format.space_after = Pt(3)


def set_table_geometry(table, widths, indent=TABLE_INDENT_DXA):
    total = sum(widths)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = Inches(widths[index] / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            set_cell_border(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_run_font(run, name="Calibri", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_style_font(style, name="Calibri", size=11, color="000000", bold=None):
    style.font.name = name
    style._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    style._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        style.font.bold = bold


def set_keep_with_next(paragraph, value=True):
    paragraph.paragraph_format.keep_with_next = value


def add_text_paragraph(doc, text, style=None, before=0, after=6, line=1.25, color=None, bold=False, italic=False):
    paragraph = doc.add_paragraph(style=style)
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line
    run = paragraph.add_run(text)
    set_run_font(run, size=11, color=color or "000000", bold=bold, italic=italic)
    return paragraph


def add_bullet(doc, text, level=0):
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.left_indent = Inches(0.375 + level * 0.25)
    paragraph.paragraph_format.first_line_indent = Inches(-0.188)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.25
    run = paragraph.add_run(text)
    set_run_font(run, size=11)
    return paragraph


def add_number(doc, text, level=0):
    paragraph = doc.add_paragraph(style="List Number")
    paragraph.paragraph_format.left_indent = Inches(0.375 + level * 0.25)
    paragraph.paragraph_format.first_line_indent = Inches(-0.188)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.25
    run = paragraph.add_run(text)
    set_run_font(run, size=11)
    return paragraph


def add_heading(doc, text, level=1):
    paragraph = doc.add_paragraph(style=f"Heading {level}")
    run = paragraph.add_run(text)
    set_keep_with_next(paragraph)
    return paragraph


def add_code_block(doc, text, fill=LIGHT_GRAY):
    paragraph = doc.add_paragraph()
    set_paragraph_box(paragraph, fill=fill, border_color="D6DEE8", border_size="4")
    paragraph.paragraph_format.line_spacing = 1.1
    run = paragraph.add_run(text)
    set_run_font(run, name="Consolas", size=9.5, color=INK)
    return paragraph


def add_callout(doc, label, text, fill=CALLOUT, label_color=BLUE):
    paragraph = doc.add_paragraph()
    set_paragraph_box(paragraph, fill=fill, border_color="C9D6E5", border_size="8")
    paragraph.paragraph_format.line_spacing = 1.2
    label_run = paragraph.add_run(label + "：")
    set_run_font(label_run, size=11, color=label_color, bold=True)
    text_run = paragraph.add_run(text)
    set_run_font(text_run, size=11, color=INK)
    return paragraph


def add_table(doc, headers, rows, widths, header_fill=LIGHT_BLUE):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths)
    header = table.rows[0]
    repeat_table_header(header)
    for index, header_text in enumerate(headers):
        cell = header.cells[index]
        set_cell_shading(cell, header_fill)
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.1
        run = paragraph.add_run(header_text)
        set_run_font(run, size=10.5, color=INK, bold=True)
    for row_data in rows:
        row = table.add_row()
        for index, value in enumerate(row_data):
            cell = row.cells[index]
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.15
            run = paragraph.add_run(str(value))
            set_run_font(run, size=10.2, color="000000")
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_checklist(doc, items):
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.left_indent = Inches(0.375)
        paragraph.paragraph_format.first_line_indent = Inches(-0.188)
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.paragraph_format.line_spacing = 1.25
        checkbox = paragraph.add_run("□ ")
        set_run_font(checkbox, size=11, color=BLUE, bold=True)
        run = paragraph.add_run(item)
        set_run_font(run, size=11)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    set_run_font(run, size=9, color=MUTED)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)
    tail = paragraph.add_run(" 页")
    set_run_font(tail, size=9, color=MUTED)


def add_title_block(doc, section):
    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hp.paragraph_format.space_after = Pt(0)
    run = hp.add_run("RehabMind  ·  起步实施方案")
    set_run_font(run, size=9, color=MUTED, bold=True)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.paragraph_format.space_before = Pt(0)
    fp.paragraph_format.space_after = Pt(0)
    add_page_number(fp)

    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(18)
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run("REHABMIND")
    set_run_font(run, size=10, color=BLUE, bold=True)

    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(5)
    paragraph.paragraph_format.line_spacing = 1.05
    run = paragraph.add_run("起步阶段实施方案")
    set_run_font(run, size=27, color=INK, bold=True)

    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(15)
    run = paragraph.add_run("单人维护、粉丝群邀请制小范围试用版本")
    set_run_font(run, size=14, color=MUTED)

    metadata = [
        ("文档用途", "起步实施、试用流程和验收基线"),
        ("当前维护模式", "项目负责人单人维护"),
        ("试用方式", "邀请制匿名案例"),
        ("首发范围", "膝关节、踝足及当前已验证流程"),
        ("文档状态", "方案已确认，待实施"),
        ("日期", "2026年8月20日"),
    ]
    for label, value in metadata:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.1
        label_run = p.add_run(label + "：")
        set_run_font(label_run, size=10.5, color=INK, bold=True)
        value_run = p.add_run(value)
        set_run_font(value_run, size=10.5, color="000000")

    rule = doc.add_paragraph()
    rule.paragraph_format.space_before = Pt(11)
    rule.paragraph_format.space_after = Pt(14)
    p_pr = rule._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), BLUE)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def setup_document():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    set_style_font(normal, size=11, color="000000")
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        style = doc.styles[name]
        set_style_font(style, size=size, color=color, bold=True)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.15
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = doc.styles[name]
        set_style_font(style, size=11, color="000000")
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)

    add_title_block(doc, section)
    return doc, section


def build_document():
    doc, section = setup_document()

    add_callout(
        doc,
        "核心结论",
        "当前不先建设复杂的多人机构平台，而是先完成“邀请用户 → 自动保存完整案例 → 后台查看和复现 → 用户反馈 → 版本迭代”的数据闭环。",
        fill=LIGHT_BLUE,
    )

    add_heading(doc, "1. 方案定位", 1)
    add_text_paragraph(doc, "本方案适用于项目由一个人维护、准备在粉丝群中小范围传播验证的阶段。重点不是快速增加功能，而是确保每个试用案例都能被你完整读取、复现和分析。")
    add_text_paragraph(doc, "产品需要能够回答四个问题：")
    for item in [
        "用户实际填写和选择了什么。",
        "系统当时识别了什么，生成了什么评估、处理、复测和训练内容。",
        "用户在哪一步卡住、退出、跳过或认为不合适。",
        "修改后是否解决了原问题，并且有没有引入新的问题。",
    ]:
        add_bullet(doc, item)
    add_code_block(doc, "邀请用户\n  ↓\n创建匿名案例并生成案例编号\n  ↓\n主诉 → 评估 → 问题清单 → 处理与即时复测 → 训练与反馈 → 总结\n  ↓\n后台查看完整时间线与用户反馈\n  ↓\n新增测试、修改、回放、灰度发布")

    add_heading(doc, "2. 起步范围与明确边界", 1)
    add_heading(doc, "2.1 首发范围", 2)
    for item in [
        "采用邀请制，不公开注册。",
        "先开放当前已经完成验证的膝关节、踝足流程。",
        "保留现有普通用户模式、康复思路模式、双侧流程以及安全出口。",
        "继续使用当前本地规则引擎，不在本阶段引入 AI 自动决策。",
        "使用匿名案例编号关联用户反馈。",
        "初期由一个管理员查看和分析案例。",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "2.2 当前暂不做", 2)
    add_table(
        doc,
        ["暂不做的内容", "原因"],
        [
            ("公开注册和复杂账号体系", "当前重点是验证流程和数据闭环，匿名案例已能满足首轮试用。"),
            ("多机构、多团队和成员协作", "目前只有一名维护者，没有实际的组织管理需求。"),
            ("微服务拆分和实时协作", "会增加部署和排错成本，不能直接解决案例无法回收的问题。"),
            ("大量新增身体区域", "先稳定膝、踝足模块，再根据真实需求决定扩展。"),
            ("自动根据案例改写知识库", "案例只能形成修改候选，不能替代资料和人工审核。"),
            ("完整可视化知识库编辑器", "初期使用版本文件和发布流程即可。"),
        ],
        [2600, 6760],
    )

    add_heading(doc, "3. 开放试用前必须补上的能力", 1)
    add_text_paragraph(doc, "当前业务记录仍主要保存在浏览器 localStorage。项目虽然已经有 Cloudflare D1 和 Drizzle 的基础脚手架，但数据库尚未承担正式案例存储。因此，开放给其他人使用前至少需要补充以下能力：")
    for item in [
        "服务端案例存储。",
        "匿名案例编号和访问凭证。",
        "自动保存和恢复。",
        "完整事件记录。",
        "单人管理员后台。",
        "案例复现。",
        "用户反馈绑定。",
        "应用、知识库和决策规则版本记录。",
        "知识缺口收集。",
        "最低限度的访问保护、删除和导出能力。",
    ]:
        add_bullet(doc, item)
    add_callout(doc, "保存原则", "页面可以保留本地草稿用于网络中断恢复，但本地数据只能是临时缓存；后台数据才是正式案例记录。", fill=CAUTION, label_color="7A5A00")

    add_heading(doc, "4. 系统边界和模块关系", 1)
    add_code_block(doc, "浏览器页面\n  ↓\n案例访问与保存接口\n  ↓\n应用服务层\n  ↓\n现有评估、处理、复测、训练决策核心\n  ↓\nD1 / Drizzle 数据库\n  ↓\n管理员案例后台与反馈分析")
    add_text_paragraph(doc, "现有的评估、处理、复测和训练核心模块继续保持纯逻辑，不直接访问数据库，不直接判断当前是谁在操作，也不把持久化逻辑塞回页面组件。新增应用服务负责读取案例、调用核心、保存输入和输出、绑定版本以及检查阶段是否允许推进。")
    add_table(
        doc,
        ["边界", "职责"],
        [
            ("页面层", "展示问题、收集输入、提示保存状态，不自行决定是否放行流程。"),
            ("应用服务层", "读取快照、调用核心、保存结果、检查权限和流程状态。"),
            ("决策核心", "根据结构化输入计算评估计划、问题、处理、复测和训练结果。"),
            ("数据层", "保存案例、快照、事件、反馈和版本，不修改业务含义。"),
            ("后台层", "查看、筛选、复现和标记问题，不直接覆盖用户原始记录。"),
        ],
        [2200, 7160],
    )

    add_heading(doc, "5. 用户端详细流程", 1)
    add_heading(doc, "5.1 邀请进入试用", 2)
    add_number(doc, "用户通过邀请链接进入试用页面。")
    add_number(doc, "页面说明这是小范围体验版本，输出是康复思路和记录辅助，不是诊断。")
    add_number(doc, "页面说明体验过程中会保存案例流程和反馈，用户可以删除案例。")
    add_number(doc, "用户确认后，服务端创建一个匿名案例。")
    add_callout(doc, "隐私边界", "初期不强制收集真实姓名、手机号、邮箱、身份证号和详细住址；用户主诉中也不应主动填写不必要的身份信息。", fill=RISK, label_color="9B1C1C")

    add_heading(doc, "5.2 创建匿名案例", 2)
    add_text_paragraph(doc, "服务端为案例生成内部 ID、用户可读的短案例编号和随机访问凭证。页面显示：")
    add_code_block(doc, "本次案例编号：RM-8K3P2\n请复制或截图保存，后续反馈时提供这个编号即可。")
    add_text_paragraph(doc, "案例编号用于和你在群里收到的反馈建立连接。用户不需要重新描述完整过程，只需要提供案例编号和问题所在阶段。访问凭证不能使用连续数字 ID，也不能把完整 token 写入公开日志。")

    add_heading(doc, "5.3 主诉和症状信息收集", 2)
    add_text_paragraph(doc, "用户按现有流程填写哪边哪里、何时或怎么出现、什么动作不舒服、症状性质、当前表现和恢复目标。不清楚的内容可以写“不清楚”。首屏示例统一使用：")
    add_callout(doc, "填写示例", "右膝内侧下楼梯时刺痛，跑步后会出现，持续两周了，平时走路还好，想恢复到正常跑步的水平。", fill=LIGHT_GRAY, label_color=INK)
    add_text_paragraph(doc, "系统需要同时保存用户原话和解析结果：")
    add_code_block(doc, "用户原话：右膝内侧下楼梯时刺痛……\n系统解析：右侧 / 膝内侧 / 下楼梯 / 跑步后出现 / 两周 / 跑步目标")
    add_text_paragraph(doc, "系统解析不能覆盖用户原话。无法识别的内容保留原文，并标记为待确认。")

    add_heading(doc, "5.4 关键确认", 2)
    for item in [
        "用户确认主要大部位、左右侧或双侧、具体不适位置、当前主要动作、目标和安全信息。",
        "解析不准确时，用户可以修改、补充或选择“不清楚”。",
        "选择答案后不自动跳转下一题，由用户点击“下一步”推进。",
        "文本解析结果只用于打开建议区域，最终确认以用户选择为准。",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "5.5 评估", 2)
    add_text_paragraph(doc, "评估继续执行当前已经确认的核心逻辑：")
    for item in [
        "主诉动作和基础功能动作合并去重。",
        "评估结果筛选后续处理、训练和复测。",
        "未完成动作不能被当作正常。",
        "没有可靠基线时不能生成普通复测资格。",
        "双侧针对性评估未完成时不能进入正常训练。",
        "不确定或证据不足时进入补充检查、低负荷活动或保存退出。",
    ]:
        add_bullet(doc, item)
    add_text_paragraph(doc, "每个评估项目需要保存项目 ID、动作名称、操作结果、完成状态、左右侧、活动度或控制结果、不适位置和分数、是否跳过以及无法完成原因。")

    add_heading(doc, "5.6 问题清单", 2)
    add_text_paragraph(doc, "评估完成后，系统将结果整理成问题清单，并说明每个问题是否有评估依据、是否需要处理、训练、复测或补充信息。问题清单是后续流程的中间结果，不是诊断结论。")
    add_callout(doc, "核心限制", "不能让“任何功能动作异常”自动扩散到所有训练。只有与当前目标、评估证据和能力链相关的问题，才进入对应下游流程。", fill=CAUTION, label_color="7A5A00")

    add_heading(doc, "5.7 处理与即时复测", 2)
    for item in [
        "显示当前处理候选和适用范围。",
        "记录用户实际执行、跳过或无法执行的处理。",
        "保存执行侧别和处理状态。",
        "按照当前规则进行即时复测。",
        "保存处理前、处理后、症状变化和活动表现变化。",
        "处理加重或活动表现变差时进入停止或重新评估出口。",
    ]:
        add_number(doc, item)
    add_text_paragraph(doc, "处理结果不能只看分数。例如疼痛 5→2，但下楼活动表现变差，必须保存为混合结果，不能标记为单纯改善。")
    add_text_paragraph(doc, "双侧处理时，同一处理目标可以合并为一张处理卡，但先处理优先侧，再处理另一侧；两侧的执行记录和复测结果独立保存，没有处理指征的一侧不强行处理。")

    add_heading(doc, "5.8 训练与反馈", 2)
    add_text_paragraph(doc, "进入训练前检查评估是否完成、双侧评估是否完成、是否存在安全停止信号、是否有可靠训练依据以及是否需要低负荷限制。")
    add_text_paragraph(doc, "训练方案可以先保存，但只有用户为实际执行的训练动作填写反馈，才算本次训练完成。每个动作记录是否执行、完成情况、是否不适、是否动作变形、是否需要退阶和是否训练后加重。")
    add_callout(doc, "跳过训练", "用户跳过训练时，可以保存“方案已保存、未执行”，但不能生成训练改善或训练进阶结论。", fill=LIGHT_GRAY, label_color=INK)

    add_heading(doc, "5.9 总结和用户反馈", 2)
    add_text_paragraph(doc, "总结页显示本次主诉、评估完成情况、主要问题、实际处理、复测变化、训练执行情况、下一次建议关注内容和案例编号。")
    add_table(
        doc,
        ["反馈选项", "用途"],
        [
            ("看不懂 / 不会做", "发现普通用户文案、说明或操作门槛问题。"),
            ("没有合适选项", "发现评估、动作或部位覆盖缺口。"),
            ("评估不合适", "发现评估项目本身不适合或与主诉不匹配。"),
            ("处理不合适", "发现候选处理重复、过多或缺少必要条件。"),
            ("复测不符合感受", "发现复测资格、复测条件或结果解释问题。"),
            ("训练不合适", "发现训练门控、负荷或反馈逻辑问题。"),
            ("页面出错 / 其他", "收集运行时错误和无法预先分类的问题。"),
        ],
        [3000, 6360],
    )

    add_heading(doc, "6. 服务端保存和数据结构", 1)
    add_heading(doc, "6.1 当前快照", 2)
    add_text_paragraph(doc, "快照用于快速恢复页面，保存案例在某一时刻的完整状态，包括当前阶段、主诉、评估、问题清单、处理队列、复测、训练计划、训练反馈、未完成项和版本信息。快照需要有递增 revision，避免重复请求覆盖较新的状态。")
    add_heading(doc, "6.2 事件记录", 2)
    add_text_paragraph(doc, "事件记录用于还原用户完整过程，原则上只追加，不直接覆盖。建议至少记录：")
    add_code_block(doc, "case_created\nconsent_confirmed\nintake_saved\nintake_confirmed\nassessment_answered\nassessment_skipped\nassessment_completed\nfinding_generated\ntreatment_started\ntreatment_skipped\ntreatment_retested\ntraining_plan_saved\ntraining_feedback_saved\ntraining_completed\nsession_saved\nfeedback_submitted\ncase_deleted")
    add_text_paragraph(doc, "每条事件保存案例 ID、事件类型、发生时间、事件顺序号、事件数据、操作来源、应用版本、知识库版本和决策规则版本。发现错误时追加修正事件，并保留原始事件。")

    add_heading(doc, "6.3 初期建议的数据表", 2)
    add_table(
        doc,
        ["数据表", "职责"],
        [
            ("pilot_cases", "匿名案例基本信息、访问凭证哈希、状态和创建时间。"),
            ("case_snapshots", "可恢复的当前快照和 revision。"),
            ("case_events", "完整事件时间线和原始操作记录。"),
            ("case_feedback", "用户反馈、所在阶段和关联事件。"),
            ("app_releases", "应用版本和发布信息。"),
            ("knowledge_releases", "知识库发布版本。"),
            ("decision_releases", "决策规则版本。"),
            ("admin_notes", "管理员分析备注，不修改用户原始记录。"),
            ("knowledge_gap_candidates", "未覆盖动作、部位和需求的候选记录。"),
        ],
        [2800, 6560],
    )

    add_heading(doc, "7. 后台查看和复现流程", 1)
    add_heading(doc, "7.1 案例列表", 2)
    add_text_paragraph(doc, "后台初期只服务项目负责人，支持按案例编号、创建时间、当前阶段、是否结束、是否安全停止、是否处理或训练加重、是否双侧、使用版本、是否有反馈和是否存在知识缺口筛选。")
    add_heading(doc, "7.2 案例详情", 2)
    add_code_block(doc, "用户原话\n  ↓\n系统解析\n  ↓\n用户确认\n  ↓\n评估输入\n  ↓\n评估结果\n  ↓\n问题清单\n  ↓\n处理记录\n  ↓\n即时复测\n  ↓\n训练反馈\n  ↓\n最终总结")
    add_text_paragraph(doc, "详情页必须区分用户填写、用户修改、系统推断、系统生成和管理员备注。管理员备注不能修改用户原始记录，只作为独立分析信息保存。")
    add_heading(doc, "7.3 案例复现", 2)
    add_text_paragraph(doc, "后台默认按照案例保存的原始版本复现：")
    add_code_block(doc, "应用版本：0.x.x\n知识库版本：knee-ankle-x.x\n决策规则版本：decision-x.x")
    add_text_paragraph(doc, "如果需要用当前版本重新计算，必须生成独立的对比结果，包含原始结果、当前版本结果、变化项目和变化原因，不能覆盖历史结果。")

    add_heading(doc, "8. 知识缺口与扩展流程", 1)
    add_text_paragraph(doc, "案例数据主要用于优化当前知识库内部的流程和决策，但如果只保存用户选择的现有选项，就会形成封闭范围偏差。因此系统必须额外保存未被覆盖的信息：")
    for item in [
        "用户原始主诉。",
        "自定义动作和自定义部位。",
        "“没有合适选项”和“无法归类”。",
        "用户认为系统输出与自身情况不符的反馈。",
        "用户认为不相关的评估或处理。",
        "专业人员的补充备注。",
    ]:
        add_bullet(doc, item)
    add_text_paragraph(doc, "后台形成独立的知识缺口候选，不直接写入正式知识库。")
    add_code_block(doc, "待观察\n  ↓\n已确认是同义表达\n  ↓\n需要资料审核\n  ↓\n设计新增条目\n  ↓\n已加入测试\n  ↓\n已发布 / 暂不纳入")
    add_callout(doc, "边界", "案例数据可以发现新增方向，但不能单独证明新的评估、处理或训练内容正确；新知识仍需结合可靠资料、专业审核、历史案例回放和小范围试用。", fill=CAUTION, label_color="7A5A00")

    add_heading(doc, "9. 知识库和规则版本迭代", 1)
    add_heading(doc, "9.1 稳定条目 ID", 2)
    add_text_paragraph(doc, "知识条目使用稳定 ID，名称可以修改，ID 不随文案变化：")
    add_code_block(doc, "assessment.function.knee.step-up\nfinding.knee.step-up.control-deficit\ntreatment.knee.posterior-chain.control\nretest.function.knee.step-up\ntraining.knee.step-up.control")
    add_heading(doc, "9.2 发布版本不可变", 2)
    add_text_paragraph(doc, "已发布的知识库版本不直接编辑。例如从 knee-ankle-0.4 创建 knee-ankle-0.4.1。新案例使用新版本，旧案例继续绑定旧版本。")
    add_heading(doc, "9.3 修改类型和验证要求", 2)
    add_table(
        doc,
        ["修改类型", "影响", "最低验证"],
        [
            ("普通文案", "通常不改变决策", "页面、移动端和理解度检查。"),
            ("评估项目", "可能改变检查队列和下游内容", "检查重复、双侧、处理、复测和训练放行。"),
            ("关联规则", "改变评估到处理、复测或训练的路径", "历史回放、差异对比、自动化测试和浏览器走读。"),
            ("安全规则", "改变停止、转介或低负荷出口", "代码变更、专门测试、人工审核和灰度发布。"),
        ],
        [2200, 3000, 4160],
    )
    add_text_paragraph(doc, "每个案例都绑定应用版本、知识库版本和决策规则版本。版本升级后不能悄悄改变旧案例的历史结论。")

    add_heading(doc, "10. 单人维护下的迭代流程", 1)
    add_code_block(doc, "后台收集案例\n  ↓\n聚合相似反馈\n  ↓\n判断问题类别\n  ↓\n形成修改假设\n  ↓\n新增可复现测试场景\n  ↓\n修改页面、流程、知识库或规则\n  ↓\n回放旧案例并比较新旧结果\n  ↓\n小范围发布\n  ↓\n观察反馈并决定保留、回滚或继续调整")
    add_table(
        doc,
        ["优先级", "问题类型"],
        [
            ("P0", "安全问题、数据丢失、串案、无法恢复。"),
            ("P1", "核心决策错误、评估与处理或复测不匹配。"),
            ("P2", "用户无法理解、无法完成、流程过长或候选过多。"),
            ("P3", "文案、页面视觉和一般体验问题。"),
            ("P4", "新功能、新部位和非关键扩展。"),
        ],
        [1600, 7760],
    )
    add_text_paragraph(doc, "每次迭代至少新增一个可复现测试场景。核心决策变化不能只靠手动看页面，必须同步加入自动化测试和真实浏览器走读。")

    add_heading(doc, "11. 试用阶段观察指标", 1)
    add_text_paragraph(doc, "以下指标用于产品流程和知识库迭代，不直接作为临床疗效结论：")
    for item in [
        "每个阶段的完成率和退出率。",
        "评估项目的跳过率、无法完成率和暂不判断比例。",
        "没有合适选项、自定义动作和自定义部位的比例。",
        "处理候选被跳过、复测未完成和训练反馈缺失的比例。",
        "用户反馈集中出现的步骤和选项。",
        "处理或训练加重记录。",
        "知识缺口候选的重复出现频率。",
        "案例保存失败、恢复失败和版本复现失败次数。",
        "新旧版本产生不同输出的案例数量。",
    ]:
        add_bullet(doc, item)
    add_callout(doc, "前期做法", "在案例数量还不多时，每周人工查看案例比一开始建设复杂报表更有价值。先把问题看懂，再决定需要哪些统计。", fill=LIGHT_BLUE)

    add_heading(doc, "12. 安全和隐私底线", 1)
    add_checklist(doc, [
        "只通过邀请链接开放试用。",
        "管理员入口受保护。",
        "案例编号使用随机值，不使用连续数字。",
        "不在公开日志记录完整主诉。",
        "不收集不必要的身份信息。",
        "用户可以删除案例。",
        "管理员操作留痕。",
        "案例和反馈不能被其他用户搜索。",
        "试用数据明确标记为体验数据。",
        "页面不把输出描述为诊断或确定性治疗结论。",
    ])

    add_heading(doc, "13. 实施顺序和验收标准", 1)
    add_heading(doc, "13.1 实施顺序", 2)
    add_table(
        doc,
        ["阶段", "实施内容", "阶段目标"],
        [
            ("第一步", "抽取案例保存接口，建立本地和服务端实现边界。", "页面和决策核心不直接依赖数据库。"),
            ("第二步", "D1 建立案例、快照、事件和反馈表，接入匿名案例编号。", "案例可保存、恢复和追踪。"),
            ("第三步", "建设单人管理员后台。", "可以查看案例列表、详情和版本。"),
            ("第四步", "加入案例复现、旧新版本对比和知识缺口候选。", "可以定位问题并形成迭代输入。"),
            ("第五步", "邀请少量粉丝试用，按周人工分析和小版本发布。", "验证真实用户数据闭环。"),
        ],
        [1400, 4700, 3260],
    )
    add_heading(doc, "13.2 起步版必须满足", 2)
    add_checklist(doc, [
        "用户刷新后不会丢失已经保存的案例。",
        "项目负责人可以从后台看到完整流程。",
        "用户反馈可以直接关联具体案例和步骤。",
        "旧案例可以按原始版本复现。",
        "用户原话、系统解析和系统推断没有混淆。",
        "安全出口不会因前端状态异常而被绕过。",
        "当前自动化测试、类型检查、构建、Lint 和关键浏览器走读保持通过。",
    ])

    add_heading(doc, "14. 起步阶段的最终判断", 1)
    add_callout(doc, "起步重点", "当前最值得优先完成的不是增加更多评估或处理条目，而是建立可靠的数据闭环。只有用户使用、系统记录、后台还原、用户反馈、问题定位和修改回放能够连起来，案例数据才真正能够服务于内部流程优化、决策修正和知识库扩展候选发现。", fill=LIGHT_BLUE)
    add_text_paragraph(doc, "建议的第一批实施任务为：服务端案例存储、匿名案例编号、关键节点自动保存、单人管理员后台和案例复现。完成这五项后，再开放粉丝群小范围试用。", bold=True, color=INK)

    doc.core_properties.title = "RehabMind 起步阶段实施方案"
    doc.core_properties.subject = "单人维护、粉丝群邀请制小范围试用的产品流程、数据闭环和版本迭代方案"
    doc.core_properties.author = "RehabMind"
    doc.core_properties.keywords = "RehabMind, pilot, case, knowledge base, versioning"
    doc.core_properties.comments = "起步实施方案"
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build_document()
